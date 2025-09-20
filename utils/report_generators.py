# utils/report_generator.py
import io
import base64
import tempfile
import pandas as pd
import numpy as np
from datetime import datetime
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.platypus import (
    SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak, Image, KeepTogether
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image as PILImage
import traceback

# NOTE: This module focuses on producing a clear, embeddable PDF/DOCX.
# For heavy-duty PDF features (true TOC with page numbers, appending external PDFs),
# consider using PyPDF2 / pikepdf to merge files post-generation.

def _safe_bytes_to_pil(image_bytes):
    """Convert raw bytes to a PIL Image or return None on failure."""
    if not image_bytes:
        return None
    try:
        return PILImage.open(io.BytesIO(image_bytes))
    except Exception:
        return None

def generate_pdf_report(report_data):
    """Generate a comprehensive PDF report with the analysis results."""
    try:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=48, leftMargin=48,
            topMargin=48, bottomMargin=48
        )

        styles = getSampleStyleSheet()
        elements = []

        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=22,
            spaceAfter=18,
            textColor=colors.HexColor('#2E86AB'),
            alignment=1  # Center aligned
        )

        heading_style = ParagraphStyle(
            'Heading1',
            parent=styles['Heading1'],
            fontSize=14,
            spaceAfter=8,
            textColor=colors.HexColor('#2E86AB'),
            spaceBefore=10
        )

        normal = styles['Normal']

        # --- Cover Page ---
        if report_data.get('include_cover', True):
            # If logo provided, embed it at top
            logo_bytes = report_data.get('logo')
            pil_logo = _safe_bytes_to_pil(logo_bytes)
            if pil_logo:
                # Resize for page width
                max_width = 4 * inch
                w, h = pil_logo.size
                ratio = min(max_width / w, 1.0)
                img_buffer = io.BytesIO()
                pil_logo = pil_logo.convert("RGBA")
                pil_logo.save(img_buffer, format='PNG')
                img_buffer.seek(0)
                reportlab_image = Image(ImageReader(img_buffer), width=w * ratio, height=h * ratio)
                # center it
                reportlab_image.hAlign = 'CENTER'
                elements.append(reportlab_image)
                elements.append(Spacer(1, 12))

            elements.append(Paragraph(report_data.get('title', 'Data Analysis Report'), title_style))
            meta_text = f"<b>Prepared by:</b> {report_data.get('author','')}" \
                        f"&nbsp;&nbsp;&nbsp;&nbsp;<b>Organization:</b> {report_data.get('company','')}" \
                        f"<br/><b>Date:</b> {report_data.get('date','')}"
            elements.append(Paragraph(meta_text, normal))
            elements.append(Spacer(1, 18))

            if report_data.get('custom_notes'):
                elements.append(Paragraph("<b>Notes:</b>", styles['Heading4']))
                elements.append(Paragraph(report_data.get('custom_notes', ''), normal))
                elements.append(Spacer(1, 12))

            # Horizontal line
            hr = Table([['']], colWidths=[16*cm], style=[('LINEABOVE', (0,0), (-1,-1), 1, colors.grey)])
            elements.append(hr)
            elements.append(PageBreak())

        # --- Table of Contents placeholder (simple list of sections) ---
        if report_data.get('include_toc', True):
            elements.append(Paragraph("Table of Contents", title_style))
            toc_lines = []
            # Build a list of included sections
            sections = []
            if report_data.get('include_exec_summary', True):
                sections.append("Executive Summary")
            if report_data.get('include_data_overview', True):
                sections.append("Data Overview")
            sections.append("Data Sample")
            if report_data.get('kpis'):
                sections.append("Key Performance Indicators")
            if report_data.get('include_statistics', True):
                sections.append("Statistical Summary")
            if report_data.get('include_missing_analysis', True):
                sections.append("Missing Values Analysis")
            if report_data.get('include_data_dict', False):
                sections.append("Data Dictionary")
            if report_data.get('processing_steps'):
                sections.append("Processing History")
            if report_data.get('include_recommendations', True):
                sections.append("Recommendations")
            if report_data.get('include_chart_gallery', False):
                sections.append("Chart Gallery (selected charts)")
            if report_data.get('appendix_file'):
                sections.append("Appendix (uploaded separately)")

            for i, s in enumerate(sections, start=1):
                elements.append(Paragraph(f"{i}. {s}", normal))
            elements.append(PageBreak())

        # --- Executive Summary ---
        if report_data.get('include_exec_summary', True):
            elements.append(Paragraph("Executive Summary", heading_style))
            summary_text = report_data.get('auto_summary') or (
                "This report provides an analysis of the dataset, including a data overview, "
                "statistical summaries, missing value analysis, and recommendations."
            )
            elements.append(Paragraph(summary_text, normal))
            elements.append(Spacer(1, 8))

        # --- Data Overview ---
        if report_data.get('include_data_overview', True):
            elements.append(Paragraph("Data Overview", heading_style))
            df = report_data['dataframe']
            overview_data = [
                ['Metric', 'Value'],
                ['Total Rows', f"{df.shape[0]:,}"],
                ['Total Columns', f"{df.shape[1]:,}"],
                ['Numeric Columns', f"{len(df.select_dtypes(include=np.number).columns):,}"],
                ['Categorical Columns', f"{len(df.select_dtypes(exclude=np.number).columns):,}"],
                ['Total Missing Values', f"{df.isna().sum().sum():,}"],
                ['Total Duplicate Rows', f"{df.duplicated().sum():,}"]
            ]
            overview_table = Table(overview_data, colWidths=[4*inch, 2.2*inch])
            overview_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2E86AB')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#F8F9FA')),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ]))
            elements.append(overview_table)
            elements.append(Spacer(1, 12))

        # --- Data Sample (first 5 rows) ---
        elements.append(Paragraph("Data Sample (First 5 Rows)", heading_style))
        df_sample = report_data['dataframe'].head(5)
        # Convert DataFrame to table data with string conversion
        table_data = [list(map(str, df_sample.columns.tolist()))]
        for _, row in df_sample.iterrows():
            table_data.append([("" if pd.isna(v) else str(v)) for v in row.tolist()])

        sample_table = Table(table_data, repeatRows=1)
        sample_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2E86AB')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 9),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 6),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#FFFFFF')),
            ('GRID', (0, 0), (-1, -1), 0.3, colors.grey),
        ]))
        elements.append(sample_table)
        elements.append(Spacer(1, 12))

        # --- KPIs ---
        if report_data.get('kpis'):
            elements.append(Paragraph("Key Performance Indicators", heading_style))
            kpi_data = [['KPI Name', 'Value']]
            for k in report_data['kpis']:
                value = k.get('value')
                if k.get('type') == 'number' and isinstance(value, (int, float, np.number)):
                    value = f"{float(value):,.2f}"
                kpi_data.append([k.get('name', ''), str(value)])
            kpi_table = Table(kpi_data, colWidths=[4*inch, 2.2*inch])
            kpi_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#A23B72')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ]))
            elements.append(kpi_table)
            elements.append(Spacer(1, 12))

        # --- Statistical Summary ---
        if report_data.get('include_statistics', True):
            numeric_cols = report_data['dataframe'].select_dtypes(include=np.number).columns.tolist()
            if numeric_cols:
                elements.append(Paragraph("Statistical Summary", heading_style))
                stats_df = report_data['dataframe'][numeric_cols].describe().T
                stats_df['variance'] = report_data['dataframe'][numeric_cols].var()
                # Build table header
                header = ['Column', 'Count', 'Mean', 'Std', 'Min', '25%', '50%', '75%', 'Max', 'Variance']
                stats_table_data = [header]
                for col in stats_df.index:
                    row = [
                        col,
                        f"{int(stats_df.loc[col,'count']):,}",
                        f"{stats_df.loc[col,'mean']:.2f}",
                        f"{stats_df.loc[col,'std']:.2f}",
                        f"{stats_df.loc[col,'min']:.2f}",
                        f"{stats_df.loc[col,'25%']:.2f}",
                        f"{stats_df.loc[col,'50%']:.2f}",
                        f"{stats_df.loc[col,'75%']:.2f}",
                        f"{stats_df.loc[col,'max']:.2f}",
                        f"{stats_df.loc[col,'variance']:.2f}"
                    ]
                    stats_table_data.append(row)

                stats_table = Table(stats_table_data, repeatRows=1)
                stats_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2E86AB')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 8),
                    ('GRID', (0, 0), (-1, -1), 0.3, colors.grey),
                ]))
                elements.append(stats_table)
                elements.append(Spacer(1, 12))

        # --- Missing values analysis ---
        if report_data.get('include_missing_analysis', True):
            mv = report_data['dataframe'].isna().sum()
            mv = mv[mv > 0].sort_values(ascending=False)
            if not mv.empty:
                elements.append(Paragraph("Missing Values Analysis", heading_style))
                missing_data = [['Column', 'Missing Count', 'Missing %']]
                for col, cnt in mv.items():
                    pct = (cnt / len(report_data['dataframe'])) * 100
                    missing_data.append([col, f"{cnt:,}", f"{pct:.2f}%"])
                missing_table = Table(missing_data, colWidths=[6*cm, 3*cm, 3*cm])
                missing_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#F18F01')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('GRID', (0, 0), (-1, -1), 0.3, colors.grey),
                ]))
                elements.append(missing_table)
                elements.append(Spacer(1, 12))

        # --- Data Dictionary ---
        if report_data.get('include_data_dict', False):
            elements.append(Paragraph("Data Dictionary", heading_style))
            df = report_data['dataframe']
            dict_df = pd.DataFrame({
                "Column": df.columns.astype(str),
                "Data Type": df.dtypes.astype(str),
                "Missing %": (df.isna().mean() * 100).round(2).astype(str) + '%',
                "Unique Values": df.nunique().astype(str)
            })
            # Build table
            dict_table_data = [list(dict_df.columns)]
            for _, r in dict_df.iterrows():
                dict_table_data.append([r['Column'], r['Data Type'], r['Missing %'], r['Unique Values']])
            dict_table = Table(dict_table_data, repeatRows=1, colWidths=[6*cm, 3.5*cm, 2.5*cm, 2.5*cm])
            dict_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2E86AB')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('GRID', (0, 0), (-1, -1), 0.3, colors.grey),
            ]))
            elements.append(dict_table)
            elements.append(Spacer(1, 12))

        # --- Processing Steps ---
        if report_data.get('processing_steps'):
            elements.append(Paragraph("Processing History", heading_style))
            for i, step in enumerate(report_data['processing_steps'], start=1):
                elements.append(Paragraph(f"{i}. {step}", normal))
                elements.append(Spacer(1, 4))
            elements.append(Spacer(1, 8))

        # --- Recommendations ---
        if report_data.get('include_recommendations', True):
            elements.append(Paragraph("Recommendations", heading_style))
            if report_data.get('recommendations'):
                for rec in report_data.get('recommendations'):
                    elements.append(Paragraph(f"- {rec}", normal))
            else:
                default_recs = [
                    "Consider imputing missing values for columns with low missing percentages.",
                    "Investigate or remove columns with high missing percentages.",
                    "Standardize numeric columns if using distance-based models.",
                    "Encode categorical variables appropriately for modeling.",
                ]
                for rec in default_recs:
                    elements.append(Paragraph(f"- {rec}", normal))
            elements.append(Spacer(1, 10))

        # --- Chart gallery placeholder ---
        if report_data.get('include_chart_gallery', False):
            elements.append(Paragraph("Chart Gallery", heading_style))
            elements.append(Paragraph("Selected charts (if any) would be embedded here. "
                                      "For large charts or many images consider zipping them separately.", normal))
            elements.append(Spacer(1, 12))

        # --- Appendix note ---
        if report_data.get('appendix_file'):
            elements.append(PageBreak())
            elements.append(Paragraph("Appendix", title_style))
            elements.append(Paragraph(
                "An appendix file was uploaded with this report. "
                "Large or complex appendix files are not embedded directly to avoid file corruption; "
                "please download the appendix separately from the application interface.",
                normal
            ))
            elements.append(Spacer(1, 12))

        # --- Footer with page numbers and optional footer text ---
        footer_text = report_data.get('footer', '')

        def _add_page_number(canvas_obj, doc_obj):
            canvas_obj.saveState()
            canvas_obj.setFont('Helvetica', 9)
            page_num = canvas_obj.getPageNumber()
            canvas_obj.drawCentredString(A4[0] / 2.0, 0.75 * cm, f"Page {page_num}")
            if footer_text:
                canvas_obj.setFont('Helvetica', 8)
                canvas_obj.drawRightString(A4[0] - 48, 0.75 * cm, footer_text)
            canvas_obj.restoreState()

        doc.build(elements, onFirstPage=_add_page_number, onLaterPages=_add_page_number)
        buffer.seek(0)
        return buffer.getvalue()
    except Exception as e:
        # In case of unexpected error, return a small PDF explaining the error to help debugging
        err_buffer = io.BytesIO()
        doc = SimpleDocTemplate(err_buffer, pagesize=A4)
        styles = getSampleStyleSheet()
        elems = [Paragraph("Error generating PDF report", styles['Title']),
                 Spacer(1, 12),
                 Paragraph(str(e), styles['Normal']),
                 Spacer(1, 6),
                 Paragraph(traceback.format_exc(), styles['Code'] if 'Code' in styles else styles['Normal'])]
        doc.build(elems)
        err_buffer.seek(0)
        return err_buffer.getvalue()

def generate_docx_report(report_data):
    """Generate a comprehensive DOCX report with the analysis results."""
    try:
        doc = Document()

        # Title / Cover
        if report_data.get('include_cover', True):
            # Logo if present
            logo_bytes = report_data.get('logo')
            if logo_bytes:
                try:
                    image = _safe_bytes_to_pil(logo_bytes)
                    if image:
                        # Save temporary resized copy to embed
                        tmp = io.BytesIO()
                        max_width_px = 800
                        w, h = image.size
                        if w > max_width_px:
                            ratio = max_width_px / w
                            image = image.resize((int(w * ratio), int(h * ratio)))
                        image.save(tmp, format='PNG')
                        tmp.seek(0)
                        doc.add_picture(tmp, width=Inches(2.5))
                        last_par = doc.paragraphs[-1]
                        last_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    pass

            title = doc.add_heading(report_data.get('title', 'Data Analysis Report'), level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            meta = doc.add_paragraph()
            meta_al = meta.alignment
            meta.add_run("Prepared by: ").bold = True
            meta.add_run(f"{report_data.get('author','')}\n")
            meta.add_run("Organization: ").bold = True
            meta.add_run(f"{report_data.get('company','')}\n")
            meta.add_run("Date: ").bold = True
            meta.add_run(f"{report_data.get('date','')}")
            meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if report_data.get('custom_notes'):
                doc.add_paragraph()
                doc.add_heading('Notes', level=2)
                doc.add_paragraph(report_data.get('custom_notes'))

            doc.add_page_break()

        # --- Simple Table of Contents (list) ---
        if report_data.get('include_toc', True):
            doc.add_heading('Table of Contents', level=1)
            sections = []
            if report_data.get('include_exec_summary', True):
                sections.append("Executive Summary")
            if report_data.get('include_data_overview', True):
                sections.append("Data Overview")
            sections.append("Data Sample")
            if report_data.get('kpis'):
                sections.append("Key Performance Indicators")
            if report_data.get('include_statistics', True):
                sections.append("Statistical Summary")
            if report_data.get('include_missing_analysis', True):
                sections.append("Missing Values Analysis")
            if report_data.get('include_data_dict', False):
                sections.append("Data Dictionary")
            if report_data.get('processing_steps'):
                sections.append("Processing History")
            if report_data.get('include_recommendations', True):
                sections.append("Recommendations")
            if report_data.get('include_chart_gallery', False):
                sections.append("Chart Gallery")
            if report_data.get('appendix_file'):
                sections.append("Appendix")
            for i, s in enumerate(sections, start=1):
                p = doc.add_paragraph(f"{i}. {s}")
            doc.add_page_break()

        # Executive Summary
        if report_data.get('include_exec_summary', True):
            doc.add_heading('Executive Summary', level=1)
            doc.add_paragraph(report_data.get('auto_summary') or
                              "This report provides an overview and insights derived from the dataset.")
        
        # Data Overview
        if report_data.get('include_data_overview', True):
            doc.add_heading('Data Overview', level=1)
            df = report_data['dataframe']
            table = doc.add_table(rows=7, cols=2)
            table.style = 'Table Grid'
            headers = ['Metric', 'Value']
            for j, h in enumerate(headers):
                table.cell(0, j).text = h
            metrics = [
                ('Total Rows', f"{df.shape[0]:,}"),
                ('Total Columns', f"{df.shape[1]:,}"),
                ('Numeric Columns', f"{len(df.select_dtypes(include=np.number).columns):,}"),
                ('Categorical Columns', f"{len(df.select_dtypes(exclude=np.number).columns):,}"),
                ('Total Missing Values', f"{df.isna().sum().sum():,}"),
                ('Total Duplicate Rows', f"{df.duplicated().sum():,}")
            ]
            for i, (metric, value) in enumerate(metrics, start=1):
                table.cell(i, 0).text = metric
                table.cell(i, 1).text = value

        # Data Sample
        doc.add_heading('Data Sample (First 5 Rows)', level=1)
        df_sample = report_data['dataframe'].head(5)
        table = doc.add_table(rows=df_sample.shape[0] + 1, cols=df_sample.shape[1])
        table.style = 'Table Grid'
        for j, col in enumerate(df_sample.columns):
            table.cell(0, j).text = str(col)
        for i, (_, row) in enumerate(df_sample.iterrows(), start=1):
            for j, val in enumerate(row):
                table.cell(i, j).text = "" if pd.isna(val) else str(val)

        # KPIs
        if report_data.get('kpis'):
            doc.add_heading('Key Performance Indicators', level=1)
            table = doc.add_table(rows=len(report_data['kpis']) + 1, cols=2)
            table.style = 'Table Grid'
            table.cell(0, 0).text = 'KPI Name'
            table.cell(0, 1).text = 'Value'
            for i, k in enumerate(report_data['kpis'], start=1):
                v = k.get('value')
                if k.get('type') == 'number' and isinstance(v, (int, float, np.number)):
                    v = f"{float(v):,.2f}"
                table.cell(i, 0).text = str(k.get('name', ''))
                table.cell(i, 1).text = str(v)

        # Statistical Summary
        if report_data.get('include_statistics', True):
            numeric_cols = report_data['dataframe'].select_dtypes(include=np.number).columns.tolist()
            if numeric_cols:
                doc.add_heading('Statistical Summary', level=1)
                stats_df = report_data['dataframe'][numeric_cols].describe().T
                stats_df['variance'] = report_data['dataframe'][numeric_cols].var()
                # Use a simplified set if many columns
                if len(numeric_cols) > 6:
                    cols_to_show = ['mean', 'std', 'min', 'max', 'variance']
                    header = ['Column', 'Mean', 'Std', 'Min', 'Max', 'Variance']
                else:
                    cols_to_show = ['count', 'mean', 'std', 'min', '25%', '50%', '75%', 'max', 'variance']
                    header = ['Column', 'Count', 'Mean', 'Std', 'Min', '25%', '50%', '75%', 'Max', 'Variance']

                table = doc.add_table(rows=len(stats_df.index) + 1, cols=len(header))
                table.style = 'Table Grid'
                for j, h in enumerate(header):
                    table.cell(0, j).text = h
                for i, col in enumerate(stats_df.index, start=1):
                    table.cell(i, 0).text = col
                    for j, stat in enumerate(cols_to_show, start=1):
                        val = stats_df.loc[col, stat]
                        if isinstance(val, (int, float, np.number)):
                            table.cell(i, j).text = f"{val:,.2f}"
                        else:
                            table.cell(i, j).text = str(val)

        # Missing Values
        if report_data.get('include_missing_analysis', True):
            mv = report_data['dataframe'].isna().sum()
            mv = mv[mv > 0].sort_values(ascending=False)
            if not mv.empty:
                doc.add_heading('Missing Values Analysis', level=1)
                table = doc.add_table(rows=len(mv) + 1, cols=3)
                table.style = 'Table Grid'
                table.cell(0, 0).text = 'Column'
                table.cell(0, 1).text = 'Missing Count'
                table.cell(0, 2).text = 'Missing %'
                for i, (col, cnt) in enumerate(mv.items(), start=1):
                    table.cell(i, 0).text = str(col)
                    table.cell(i, 1).text = f"{int(cnt):,}"
                    table.cell(i, 2).text = f"{(cnt / len(report_data['dataframe']) * 100):.2f}%"

        # Data Dictionary
        if report_data.get('include_data_dict', False):
            doc.add_heading('Data Dictionary', level=1)
            df = report_data['dataframe']
            dict_df = pd.DataFrame({
                "Column": df.columns.astype(str),
                "Data Type": df.dtypes.astype(str),
                "Missing %": (df.isna().mean() * 100).round(2).astype(str) + '%',
                "Unique Values": df.nunique().astype(str)
            })
            table = doc.add_table(rows=len(dict_df) + 1, cols=4)
            table.style = 'Table Grid'
            headers = list(dict_df.columns)
            for j, h in enumerate(headers):
                table.cell(0, j).text = h
            for i, r in enumerate(dict_df.itertuples(index=False), start=1):
                table.cell(i, 0).text = str(r[0])
                table.cell(i, 1).text = str(r[1])
                table.cell(i, 2).text = str(r[2])
                table.cell(i, 3).text = str(r[3])

        # Processing steps
        if report_data.get('processing_steps'):
            doc.add_heading('Processing History', level=1)
            for step in report_data['processing_steps']:
                doc.add_paragraph(step, style='List Bullet')

        # Recommendations
        if report_data.get('include_recommendations', True):
            doc.add_heading('Recommendations', level=1)
            if report_data.get('recommendations'):
                for rec in report_data['recommendations']:
                    doc.add_paragraph(rec, style='List Bullet')
            else:
                defaults = [
                    "Consider imputing missing values for columns with low missing percentages.",
                    "Investigate or remove columns with high percentages of missing values."
                ]
                for r in defaults:
                    doc.add_paragraph(r, style='List Bullet')

        # Chart gallery placeholder
        if report_data.get('include_chart_gallery', False):
            doc.add_heading('Chart Gallery', level=1)
            doc.add_paragraph("Charts (if any) would be embedded here. For many charts, consider exporting them separately.")

        # Appendix note
        if report_data.get('appendix_file'):
            doc.add_page_break()
            doc.add_heading('Appendix', level=1)
            doc.add_paragraph(
                "An appendix file was uploaded alongside this report. Large appendix files are available for download separately."
            )

        # Footer: docx does not have a direct simple footer API in python-docx across all environments;
        # we keep the metadata and leave footer handling to Word if needed.

        # Save to buffer and return bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    except Exception as e:
        # Fallback: create a small DOCX describing the error
        buf = io.BytesIO()
        doc = Document()
        doc.add_heading("Error generating DOCX report", level=1)
        doc.add_paragraph(str(e))
        doc.add_paragraph(traceback.format_exc())
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()
